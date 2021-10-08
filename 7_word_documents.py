import docx

d = docx.Document('document.docx')
d.paragraphs
p = d.paragraphs[1]
p.runs
p.runs[0].bold = True
d.add_paragraph('Hello, world!')
d.save('newdoc.docx')